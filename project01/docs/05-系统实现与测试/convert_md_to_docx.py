"""Convert markdown system document to .docx with formatting preserved."""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_FILE = r"F:\Stable\project01\docs\05-系统实现与测试\系统实现与测试文档V1.md"
DOCX_FILE = r"F:\Stable\project01\docs\05-系统实现与测试\系统实现与测试文档V1.docx"


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_styled(doc, text, level):
    """Add heading with proper Chinese font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_paragraph_styled(doc, text="", bold=False, size=None, color=None, alignment=None, space_after=None):
    """Add paragraph with styling."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_code_block(doc, code_text):
    """Add code block as formatted paragraph."""
    lines = code_text.strip().split('\n')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    for line in lines:
        if line:
            run = p.add_run(line + '\n')
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        else:
            p.add_run('\n')
    # Add shading for code block
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
    pPr.append(shd)


def parse_inline_markdown(text):
    """Parse inline markdown and return list of (text, bold, code) tuples."""
    parts = []
    # Process code first
    remaining = text
    while remaining:
        # Check for inline code
        code_match = re.search(r'`([^`]+)`', remaining)
        # Check for bold
        bold_match = re.search(r'\*\*([^*]+)\*\*', remaining)

        if not code_match and not bold_match:
            parts.append((remaining, False, False))
            break

        # Find earliest match
        candidates = []
        if code_match:
            candidates.append((code_match.start(), code_match))
        if bold_match:
            candidates.append((bold_match.start(), bold_match))
        candidates.sort(key=lambda x: x[0])

        start, match = candidates[0]
        if start > 0:
            parts.append((remaining[:start], False, False))

        is_code = match.lastgroup is None and match.re == re.compile(r'`([^`]+)`')
        # Determine if it's code or bold
        if text[match.start()] == '`':
            parts.append((match.group(1), False, True))
        else:
            parts.append((match.group(1), True, False))

        remaining = remaining[match.end():]
    return parts


def add_rich_paragraph(doc, text, size=10.5, alignment=None):
    """Add paragraph with inline formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)

    # Parse and add runs
    parts = parse_inline_markdown(text)
    if not parts:
        parts = [(text, False, False)]

    for part_text, is_bold, is_code in parts:
        run = p.add_run(part_text)
        if is_code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE0, 0x34, 0x34)
        else:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(size)
            run.bold = is_bold
    return p


def add_table_from_md(doc, lines):
    """Parse and add a markdown table."""
    # Find table boundaries
    table_rows = []
    i = 0
    while i < len(lines) and not lines[i].strip().startswith('|'):
        i += 1
    if i >= len(lines):
        return i

    # Collect all rows
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_rows.append(lines[i].strip())
        i += 1

    if len(table_rows) < 2:
        return i

    # Parse header
    header_cells = [c.strip() for c in table_rows[0].split('|')[1:-1]]
    # Skip separator row (index 1)
    # Parse data rows
    data_rows = []
    for row in table_rows[2:]:
        cells = [c.strip() for c in row.split('|')[1:-1]]
        if cells:
            data_rows.append(cells)

    if not header_cells:
        return i

    num_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "E8F0FE")

    # Data rows
    for r_idx, row_cells in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < num_cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                parts = parse_inline_markdown(cell_text)
                if not parts:
                    parts = [(cell_text, False, False)]
                for part_text, is_bold, is_code in parts:
                    run = p.add_run(part_text)
                    if is_code:
                        run.font.name = "Consolas"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0xE0, 0x34, 0x34)
                    else:
                        run.font.name = "Microsoft YaHei"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                        run.font.size = Pt(9)
                        run.bold = is_bold
                p.paragraph_format.line_spacing = 1.2

    # set alternating row colors
    for r_idx in range(len(data_rows)):
        if r_idx % 2 == 1:
            for c_idx in range(num_cols):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "F8F9FA")

    return i


def add_list_item(doc, text, level=0, ordered=False, number=1):
    """Add a list item."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)

    prefix = f"{number}. " if ordered else "• "
    if level > 0:
        prefix = "  " * level + ("- " if not ordered else f"{number}. ")
    run = p.add_run(prefix)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)

    parts = parse_inline_markdown(text)
    if not parts:
        parts = [(text, False, False)]
    for part_text, is_bold, is_code in parts:
        run = p.add_run(part_text)
        if is_code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE0, 0x34, 0x34)
        else:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10.5)
            run.bold = is_bold
    return p


def convert():
    """Main conversion function."""
    with open(MD_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    lines = content.split('\n')
    in_code_block = False
    code_buffer = []
    in_table = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip table separator row
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            i += 1
            continue

        # Table handling
        if stripped.startswith('|') and stripped.endswith('|'):
            next_idx = i + 1
            while next_idx < len(lines) and lines[next_idx].strip().startswith('|'):
                next_idx += 1
            i = add_table_from_md(doc, lines[i:])
            # Let add_table_from_md handle all table rows
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headers
        if stripped.startswith('# '):
            add_heading_styled(doc, stripped[2:], 0)
            i += 1
            continue
        elif stripped.startswith('## '):
            add_heading_styled(doc, stripped[3:], 1)
            i += 1
            continue
        elif stripped.startswith('### '):
            add_heading_styled(doc, stripped[4:], 2)
            i += 1
            continue
        elif stripped.startswith('#### '):
            add_heading_styled(doc, stripped[5:], 3)
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Unordered list
        if re.match(r'^- ', stripped) or re.match(r'^\* ', stripped):
            level = 0
            # Count leading spaces to determine nesting
            leading = len(line) - len(line.lstrip())
            if leading > 0:
                level = leading // 2
            add_list_item(doc, stripped[2:], level=level, ordered=False)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if ol_match:
            number = int(re.match(r'(\d+)', stripped).group(1))
            add_list_item(doc, ol_match.group(1), ordered=True, number=number)
            i += 1
            continue

        # Regular paragraph
        add_rich_paragraph(doc, stripped)
        i += 1

    # Add page number in footer
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Nebula Studio — 系统实现与测试文档 V1")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(DOCX_FILE)
    print(f"✅ 转换完成: {DOCX_FILE}")


if __name__ == "__main__":
    convert()
